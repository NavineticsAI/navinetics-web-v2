"""Build the stakeholder review document from the copy manifest.

    python tools/copy-docx.py [--manifest copy/copy-manifest.json] [--out copy/]

THE DOCUMENT HAS ONE JOB: a reviewer opens it, finds a sentence, and knows
where that sentence is on the website. Everything here serves that and nothing
else. The first draft explained the mechanism at length on the cover and
repeated each string's full address on every row; both were cut, because a
reviewer who has to read instructions to find a paragraph will not find it.

So: page, then section, then the text. The section headings carry the location,
which is why a row only has to say WHAT it is — "Heading", "Paragraph 2".

WHY A TABLE. The obvious layout is a small grey code with the sentence under it.
It reads well and does not survive a reviewer: the code is just another line of
text, so it gets selected with the paragraph above and deleted, and the edit can
no longer be attributed to anything. A cell is something you edit INSIDE.
Track Changes works normally in table cells.
"""
import argparse
import json
import os
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

GREY = RGBColor(0x77, 0x7B, 0x80)
INK = RGBColor(0x1A, 0x1D, 0x21)
ACTION = RGBColor(0x0B, 0x5F, 0x87)
RULE = 'D9DCE0'
BAND = 'F4F6F7'


def shade(cell, fill):
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear')
    el.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(el)


def borders(table):
    """Hairlines. A heavy grid turns 1,500 rows into a wall."""
    el = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), '4')
        e.set(qn('w:color'), RULE)
        el.append(e)
    table._tbl.tblPr.append(el)


def repeat_header(row):
    pr = row._tr.get_or_add_trPr()
    el = OxmlElement('w:tblHeader')
    el.set(qn('w:val'), 'true')
    pr.append(el)


def run(p, text, size=10, bold=False, italic=False, color=INK, font='Aptos'):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    r.font.name = font
    return r


def para(doc, space_after=6, space_before=0, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    return p


def add_table(doc, rows, note_col=True):
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    borders(t)
    hdr = t.rows[0]
    repeat_header(hdr)
    for cell, label in zip(hdr.cells, ('What it is', 'Text  —  edit here')):
        cell.paragraphs[0].text = ''
        run(cell.paragraphs[0], label, size=8, bold=True, color=GREY)
        shade(cell, BAND)

    for e in rows:
        r = t.add_row()
        left, right = r.cells

        lp = left.paragraphs[0]
        lp.paragraph_format.space_after = Pt(0)
        run(lp, e['label'], size=9, color=INK)
        code = left.add_paragraph()
        code.paragraph_format.space_after = Pt(0)
        run(code, f'⟦{e["id"]}⟧', size=7, color=GREY, font='Consolas')
        if note_col and e.get('note'):
            n = left.add_paragraph()
            n.paragraph_format.space_after = Pt(0)
            run(n, e['note'], size=7.5, italic=True, color=GREY)
        shade(left, BAND)

        # One Word paragraph per line of the original, so a line break somebody
        # put in a headline survives the round trip as a line break.
        lines = e['text'].split('\n')
        rp = right.paragraphs[0]
        rp.paragraph_format.space_after = Pt(0)
        run(rp, lines[0], size=10.5)
        for extra in lines[1:]:
            q = right.add_paragraph()
            q.paragraph_format.space_after = Pt(0)
            run(q, extra, size=10.5)

    for row in t.rows:
        row.cells[0].width = Inches(1.7)
        row.cells[1].width = Inches(5.2)
    return t


def build(manifest, out_path, generated_on):
    doc = Document()

    base = doc.styles['Normal']
    base.font.name = 'Aptos'
    base.font.size = Pt(10.5)
    base.paragraph_format.space_after = Pt(6)
    for name, size, color, bold in (('Heading 1', 19, INK, True),
                                    ('Heading 2', 13, ACTION, True),
                                    ('Heading 3', 10.5, INK, True)):
        st = doc.styles[name]
        st.font.name = 'Aptos'
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = bold

    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(0.7)
    sec.top_margin = sec.bottom_margin = Inches(0.7)

    entries = manifest['entries']
    live = [e for e in entries if e.get('rendered') is not False]
    dead = [e for e in entries if e.get('rendered') is False]
    parts = [p for p in manifest['parts'] if p['count']]

    # ── cover ────────────────────────────────────────────────────────────
    # THIS PAGE IS THE PRODUCT. Everything else is a table. Five directors and
    # a regulatory consultant will read this once, in a hurry, six months apart,
    # and nobody will explain it to them — so it is three numbered steps, in
    # order, with the thing they must do LAST given its own box. The earlier
    # draft opened with four rules and a switch and a section on searching, and
    # buried the one action that makes anything happen.
    run(para(doc, space_after=2), 'NaviNetics website', size=26, bold=True)
    run(para(doc, space_after=3), 'Every word on the site, for review', size=14, color=ACTION)
    run(para(doc, space_after=4),
        f'{generated_on}   ·   {len(live)} pieces of text', size=9, color=GREY)
    run(para(doc, space_after=16),
        'Change the wording here and it goes onto the website. Nobody has to '
        'retype anything, and you do not need any software except Word.',
        size=10.5, color=GREY)

    run(para(doc, space_after=9), 'How to use this document', size=14, bold=True, color=ACTION)

    steps = [
        ('Turn on Track Changes.',
         'Review ▸ Track Changes. Do this first, so we can see who changed what. '
         'If it is off your edits still work, but they arrive with no name on them.'),
        ('Edit the right-hand column. Only that column.',
         'The grey column on the left is the address of each piece of text — which page '
         'it is on and what it is. It is how your wording finds its way back to the '
         'website. Change it and that edit is lost.'),
        ('When you have finished, change NO to YES below, and save.',
         'Nothing reaches the website until you do. Edit and save as much as you like '
         'first — for an hour, over a week — and nothing happens.'),
    ]
    for n, (what, why) in enumerate(steps, 1):
        pp = para(doc, space_after=2)
        run(pp, f'{n}.  ', size=12, bold=True, color=ACTION)
        run(pp, what, size=12, bold=True)
        b = para(doc, space_after=9, indent=0.3)
        run(b, why, size=10, color=GREY)

    # ── the publish switch ───────────────────────────────────────────────
    # Word autosaves to OneDrive every few seconds. Without this, twenty minutes
    # of editing would publish fifty times. It is the reviewer saying "I have
    # finished" — one cell, no macro, no button, no add-in, nothing to install,
    # because a macro would run on their laptop and their laptop has no copy of
    # the website.
    sw = doc.add_table(rows=1, cols=2)
    sw.autofit = False
    borders(sw)
    left, right = sw.rows[0].cells
    lp = left.paragraphs[0]
    lp.paragraph_format.space_after = Pt(0)
    run(lp, 'Publish to the website', size=11, bold=True)
    code = left.add_paragraph()
    code.paragraph_format.space_after = Pt(0)
    run(code, '⟦PUBLISH⟧', size=7, color=GREY, font='Consolas')
    shade(left, BAND)
    rp = right.paragraphs[0]
    rp.paragraph_format.space_after = Pt(0)
    run(rp, 'NO', size=20, bold=True)
    run(rp, '        ← change to YES when you are done', size=9.5, color=GREY)
    left.width = Inches(2.2)
    right.width = Inches(4.7)
    run(para(doc, space_before=7, space_after=4),
        'Your changes appear on the website a few minutes later. Set it back to '
        'NO afterwards, so the next round is deliberate too.', size=10, color=GREY)
    # The export this document was built from. Everything the reviewer types is
    # compared against the copy AS IT WAS THEN, not as it is now — see s5 of the
    # spec. Without it a document made before a correction was published, and
    # returned after, silently puts the old wording back.
    bp = para(doc, space_after=16)
    run(bp, 'Document version ', size=8, color=GREY)
    run(bp, f'⟦BASE:{manifest["generated_from"]}⟧', size=7, color=GREY, font='Consolas')

    run(para(doc, space_after=6), 'Two things to avoid', size=13, bold=True, color=ACTION)
    for what, why in [
        ('Do not add or delete rows.',
         'To have a sentence removed, or a new one added, leave a comment instead '
         '(Review ▸ New Comment). Both need a change to the page itself, which we '
         'will make.'),
        ('Keep anything written like {1} exactly as it is.',
         'The website fills those in — a name, a number, a year. Move the words '
         'around them freely, but every {1} has to stay, in the same order.'),
    ]:
        pp = para(doc, space_after=2)
        run(pp, '•  ', size=11, bold=True, color=ACTION)
        run(pp, what, size=11, bold=True)
        b = para(doc, space_after=8, indent=0.3)
        run(b, why, size=10, color=GREY)

    run(para(doc, space_before=6, space_after=3), 'Finding what you want to change',
        size=13, bold=True, color=ACTION)
    run(para(doc, space_after=4),
        'Ctrl+F and paste in the sentence you are looking for. Or use the contents '
        'on the next page — the pages are in the same order as the website menu, '
        'and each one says its web address.', size=10, color=GREY)

    run(para(doc, space_before=10, space_after=3), 'If something cannot be done',
        size=13, bold=True, color=ACTION)
    run(para(doc, space_after=4),
        'Nothing is applied unless all of it can be. If one of your changes cannot '
        'go through — a {1} was lost, or a page changed while you had this open — '
        'then none of them do, and Shubham is told why. You will not get half of '
        'an edit.', size=10, color=GREY)

    # ── contents ─────────────────────────────────────────────────────────────
    doc.add_page_break()
    run(para(doc, space_after=10), 'Contents', size=17, bold=True)
    for group in ('Site-wide', 'Pages', 'Shared content'):
        gp = [p for p in parts if p['group'] == group]
        if not gp:
            continue
        g = para(doc, space_before=12, space_after=3)
        run(g, group, size=11, bold=True, color=ACTION)
        if group == 'Shared content':
            run(g, '   — used by more than one page', size=9, color=GREY)
        for p in gp:
            row = para(doc, space_after=1, indent=0.2)
            run(row, p['title'], size=10)
            if p.get('path'):
                run(row, f'   {p["path"]}', size=8.5, color=GREY)
            if p.get('unlisted'):
                run(row, '   not in the menus', size=8.5, italic=True, color=GREY)

    # ── the copy ─────────────────────────────────────────────────────────────
    by_part = {}
    for e in live:
        by_part.setdefault(e['part'], []).append(e)

    for group in ('Site-wide', 'Pages', 'Shared content'):
        gp = [p for p in parts if p['group'] == group and by_part.get(p['key'])]
        if not gp:
            continue
        doc.add_page_break()
        doc.add_heading(group, level=1)

        for part in gp:
            rows = by_part[part['key']]
            doc.add_heading(part['title'], level=2)

            m = para(doc, space_after=8)
            if part.get('path'):
                run(m, part['path'], size=9, color=GREY)
                if part.get('unlisted'):
                    run(m, '   ·   page is not linked from the menus', size=9, italic=True, color=GREY)
            if part.get('appears_on'):
                if part.get('path'):
                    run(m, '\n', size=9)
                run(m, 'Changing anything here changes: ', size=9, bold=True, color=GREY)
                run(m, ', '.join(part['appears_on']), size=9, color=GREY)

            # Grouped by the section a reader would see it in, in page order.
            block, title = [], rows[0].get('section', '')
            for e in rows + [None]:
                s = e.get('section', '') if e else '\x00'
                if e is None or s != title:
                    if block:
                        if title:
                            run(para(doc, space_before=9, space_after=3), title,
                                size=10.5, bold=True)
                        add_table(doc, block)
                    block, title = [], s
                if e:
                    block.append(e)

    # ── appendix ─────────────────────────────────────────────────────────────
    if dead:
        doc.add_page_break()
        doc.add_heading('Appendix — text we could not find on the live pages', level=1)
        run(para(doc, space_after=10),
            'Each of these is written into the site but did not appear when we '
            'checked every page. Either it only shows in a particular situation — '
            'a form while it is sending, a menu while it is open — or that part of '
            'the site is currently switched off. It is here for completeness. '
            'Editing it may not change anything you can see.', size=10, color=GREY)
        by_file = {}
        for e in dead:
            by_file.setdefault(e['file'], []).append(e)
        for f, rows in sorted(by_file.items()):
            run(para(doc, space_before=9, space_after=3), f.replace('src/', ''),
                size=10, bold=True, color=GREY)
            add_table(doc, rows, note_col=False)

    doc.save(out_path)
    return len(live), len(dead)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--manifest', default='copy/copy-manifest.json')
    ap.add_argument('--out', default='copy')
    ap.add_argument('--date', default='')
    ap.add_argument('--name', default='',
                    help='output filename; defaults to a dated review name')
    a = ap.parse_args()

    with open(a.manifest, encoding='utf-8') as f:
        manifest = json.load(f)

    os.makedirs(a.out, exist_ok=True)
    # The filename is the first thing a reviewer sees, in an inbox next to
    # forty other attachments. It says what it is and which round it is, so
    # that a second review does not overwrite the first in someone's folder.
    name = a.name or f'NaviNetics Website Review {a.date or "draft"}.docx'
    out_path = os.path.join(a.out, name)
    live, dead = build(manifest, out_path, a.date or 'Generated today')
    print(f'\n  {live} pieces of text, {dead} in the appendix -> {out_path}\n')
    return 0


if __name__ == '__main__':
    sys.exit(main())
